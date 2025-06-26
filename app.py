# import streamlit as st
# from transformers import BertForSequenceClassification, BertTokenizerFast
# import torch
# import pickle
# import numpy as np
# import re
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity
# from PyPDF2 import PdfReader
# from docx import Document
# from io import StringIO

# import PyPDF2
# import docx

# # === Load Model, Tokenizer, and Label Encoder for Category Prediction ===
# # @st.cache_resource
# # def load_model():
# #     model = BertForSequenceClassification.from_pretrained("bert_resume_model")
# #     tokenizer = BertTokenizerFast.from_pretrained("bert_resume_model")
# #     with open(r"C:\Users\MANISH\Desktop\Mini Project\resume_predictionApp\bert_resume_model\label_encoder.pkl", "rb") as f:
# #         le = pickle.load(f)
# #     return model, tokenizer, le
# model = BertForSequenceClassification.from_pretrained("predator279/resume-classifier-model")
# tokenizer = BertTokenizerFast.from_pretrained("predator279/resume-classifier-model")

# from huggingface_hub import hf_hub_download

# label_encoder_path = hf_hub_download(
#     repo_id="predator279/resume-classifier-model",
#     filename="label_encoder.pkl"
# )

import streamlit as st
from transformers import BertForSequenceClassification, BertTokenizerFast
import torch
import pickle
import numpy as np
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from PyPDF2 import PdfReader
from docx import Document
from io import StringIO
from huggingface_hub import hf_hub_download

import PyPDF2
import docx

# === Load Model, Tokenizer, and Label Encoder for Category Prediction ===
@st.cache_resource
def load_model():
    model = BertForSequenceClassification.from_pretrained("predator279/resume-classifier-model")
    tokenizer = BertTokenizerFast.from_pretrained("predator279/resume-classifier-model")
    label_encoder_path = hf_hub_download(
        repo_id="predator279/resume-classifier-model",
        filename="label_encoder.pkl"
    )
    with open(label_encoder_path, "rb") as f:
        le = pickle.load(f)
    return model, tokenizer, le

# ... rest of your code ...


# === Text Cleaning Function for Both Categories and Ranking ===
def clean_text(text):
    text = str(text)
    text = re.sub(r'http\S+', '', text)
    text = re.sub(r'RT|cc', '', text)
    text = re.sub(r'#\S+', '', text)
    text = re.sub(r'@\S+', '', text)
    text = re.sub(r'[^\w\s]', '', text)
    text = re.sub(r'[^\x00-\x7f]', r' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# === Function to Extract Text from Files for Ranking ===
def extract_text_from_file(file):
    if file.type == "application/pdf":
        reader = PyPDF2.PdfReader(file)
        return " ".join([page.extract_text() or '' for page in reader.pages])
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file.type == "text/plain":
        return file.read().decode("utf-8")
    else:
        return ""

# === Function to Extract Text from Files ===
def extract_text_from_pdf(uploaded_file):
    reader = PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def extract_text_from_txt(uploaded_file):
    return uploaded_file.getvalue().decode("utf-8")

# === Resume Category Prediction ===
def predict_category(text, model, tokenizer, le, device, top_k=5):
    cleaned = clean_text(text)
    max_length = 512
    inputs = tokenizer(cleaned, return_tensors="pt", truncation=True, padding=True, max_length=max_length)
    inputs = {k: v.to(device) for k, v in inputs.items()}
    
    with torch.no_grad():
        logits = model(**inputs).logits
        probs = torch.softmax(logits, dim=1).cpu().numpy()[0]
    
    top_indices = np.argsort(probs)[::-1][:top_k]
    top_labels = le.inverse_transform(top_indices)
    top_scores = probs[top_indices]
    
    return list(zip(top_labels, top_scores))

# === Resume Ranking Based on Job Description ===
def rank_resumes(job_description, uploaded_files):
    job_description_cleaned = clean_text(job_description)
    
    resume_texts = []
    resume_names = []
    similarity_scores = []
    score_breakdowns = []

    for file in uploaded_files:
        resume_raw = extract_text_from_file(file)
        resume_cleaned = clean_text(resume_raw)
        resume_texts.append(resume_cleaned)
        resume_names.append(file.name)

        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf_matrix = vectorizer.fit_transform([job_description_cleaned, resume_cleaned])
        
        score = cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0]
        similarity_scores.append(score)

        feature_names = vectorizer.get_feature_names_out()
        job_tfidf = tfidf_matrix[0].toarray()[0]
        resume_tfidf = tfidf_matrix[1].toarray()[0]

        common_terms = [
            (feature_names[i], job_tfidf[i], resume_tfidf[i])
            for i in np.where(resume_tfidf > 0)[0] if job_tfidf[i] > 0
        ]
        score_breakdowns.append(common_terms)

    ranked_indices = np.argsort(similarity_scores)[::-1]
    return resume_names, similarity_scores, score_breakdowns, ranked_indices

# === Streamlit UI ===
st.set_page_config(page_title="Resume Analyzer", layout="centered")
st.title("🔍 AI-Powered Resume Analyzer")

# Sidebar for navigation
mode = st.sidebar.radio("Choose Mode:", ["Resume Category Prediction", "Resume Ranking"])

# Category Prediction Mode
if mode == "Resume Category Prediction":
    st.subheader("Resume Category Prediction (BERT)")

    uploaded_file = st.file_uploader("Upload your resume (.pdf, .txt, .docx)", type=["pdf", "txt", "docx"])

    if uploaded_file is not None:
        model, tokenizer, le = load_model()
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        model.to(device)
        model.eval()

        if uploaded_file.type == "application/pdf":
            resume_text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            resume_text = extract_text_from_docx(uploaded_file)
        else:
            resume_text = extract_text_from_txt(uploaded_file)

        if resume_text:
            st.subheader("📄 Resume Preview")
            st.write(resume_text[:500] + "..." if len(resume_text) > 500 else resume_text)

            st.subheader("🔮 Top 5 Predicted Categories")
            top5 = predict_category(resume_text, model, tokenizer, le, device)
            for i, (label, score) in enumerate(top5, 1):
                st.write(f"**{i}. {label}** — Score: {score:.4f}")
        else:
            st.error("❌ Unsupported file type or empty content.")

# Ranking Mode
elif mode == "Resume Ranking":
    st.subheader("Resume Ranking Based on Job Description")
    
    job_description = st.text_area("📝 Paste Job Description", height=200)
    uploaded_resumes = st.file_uploader("📄 Upload Resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True)

    if st.button("🚀 Rank Resumes") and job_description and uploaded_resumes:
        resume_names, similarity_scores, score_breakdowns, ranked_indices = rank_resumes(job_description, uploaded_resumes)

        st.subheader("📊 Ranked Resumes:")
        for rank, idx in enumerate(ranked_indices):
            st.markdown(f"**{rank + 1}. {resume_names[idx]}** — Similarity Score: `{similarity_scores[idx]:.4f}`")
            with st.expander("🔍 View Matching Keywords"):
                for term, job_score, resume_score in score_breakdowns[idx]:
                    st.write(f"`{term}`: JD({job_score:.4f}) | Resume({resume_score:.4f})")
